import aiofiles
import asyncio
import json
import os
from aiopath import AsyncPath
from typing import AsyncGenerator, List, Dict, Tuple, Optional, Set, Union
from docx import Document
from docxcompose.composer import Composer
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import logging
from openai import AsyncOpenAI
from dotenv import load_dotenv

from src.elasticsearch_service import ElasticsearchService
from src.csv_export import AraDocument, EdarehoquqyDocument
from src.html_converter import get_doc_title, get_each_doc_summary, get_all_existing_ids
from src.html_export import save_html_file

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

load_dotenv()


class Loader:
    def __init__(self, root_dir: str = "search", elasticsearch_service: ElasticsearchService = None, openai_client: AsyncOpenAI = None):
        """
        Initialize a FileLoader to process JSON files from the search directory structure.
        
        Args:
            root_dir: Root directory containing keyword subdirectories
        """
        self.root_dir = AsyncPath(root_dir)
        self._failed_files = set()
        self._failed_ids: Set[str] = set()
        self._all_doc_ids: Set[str] = set()
        self._categorized_ids: Set[str] = set()
        self.elasticsearch_service = elasticsearch_service
        self.openai_client = openai_client
        
    async def get_keywords(self) -> List[str]:
        """Get all keyword directories in the root directory"""
        if not await self.root_dir.exists():
            logger.error(f"Root directory {self.root_dir} does not exist")
            return []
            
        keyword_dirs = []
        async for _dir in self.root_dir.glob("*"):
            if await _dir.is_dir() and not _dir.name.startswith('.') and _dir.name != "بدون-دسته-بندی":
                keyword_dirs.append(_dir.name)
        return keyword_dirs

    async def get_all_existing_ids(self, search_dir: str) -> List[str]:
        path = AsyncPath(search_dir)
        if not await path.exists():
            logger.warning(f"Search directory '{search_dir}' doesn't exist!")
            return []
        
        logger.info(f"Scanning for IDs in directory: {search_dir}")
        existing_ids = set()
        async for keyword_dir in path.iterdir():
            if await keyword_dir.is_dir():
                logger.info(f"Found keyword directory: {keyword_dir}")
                async for id_dir in keyword_dir.iterdir():
                    if await id_dir.is_dir():
                        # Add the directory name (id) to our set
                        logger.info(f"Found ID directory: {id_dir.name}")
                        existing_ids.add(id_dir.name)
        
        uncategorized_dir = AsyncPath(os.path.join(str(self.root_dir), "بدون-دسته-بندی"))
        if not await uncategorized_dir.exists():
            logger.info(f"Uncategorized directory '{uncategorized_dir}' doesn't exist!")
            return list(existing_ids)

        # Add uncategorized IDs
        async for id_dir in uncategorized_dir.iterdir():
            if await id_dir.is_dir():
                existing_ids.add(id_dir.name)

        logger.info(f"Total IDs found: {len(existing_ids)}")
        return list(existing_ids)


    async def get_id_dirs(self, keyword: str) -> List[AsyncPath]:
        """Get all ID directories within a keyword directory"""
        # Convert spaces to underscores to match directory naming convention
        formatted_keyword = keyword.replace('/', '_').replace(' ', '_')
        keyword_path = self.root_dir / formatted_keyword
        if not await keyword_path.exists():
            logger.error(f"Keyword directory {keyword_path} does not exist")
            return []
            
        id_dirs = []
        async for _dir in keyword_path.glob("*"):
            if await _dir.is_dir() and not _dir.name.startswith('.'):
                self._all_doc_ids.add(_dir.name)  
                id_dirs.append(_dir)
        logger.info(f"Found {len(id_dirs)} ID directories for keyword: {keyword}")
        return id_dirs  
    
    async def load_json_files(self, dir_path: AsyncPath, index_type: str = "edarehoquqy") -> AsyncGenerator[Union[EdarehoquqyDocument, AraDocument], None]:
        """Load JSON files from an ID directory and convert them to document objects based on index type"""
        try:
            json_files = []
            async for file in dir_path.glob("**/*.json"):
                json_files.append(file)
                
            if not json_files:
                logger.warning(f"No JSON files found in {dir_path}")
                return
                
            for file in json_files:
                try:
                    async with aiofiles.open(file, mode='r', encoding='utf-8') as f:
                        data = await f.read()
                        json_data = json.loads(data)
                        if index_type == "edarehoquqy":
                            logger.info(f"Loading EdarehoquqyDocument from {file}")
                            yield EdarehoquqyDocument(**json_data)
                        elif index_type == "ara":
                            yield AraDocument(**json_data)
                        else:
                            logger.warning(f"Unknown index type: {index_type}")
                            continue
                except json.JSONDecodeError:
                    logger.error(f"Failed to decode JSON from {file}")
                    self._failed_files.add(str(file))
                except Exception as e:
                    logger.error(f"Error loading file {file}: {e}")
                    self._failed_files.add(str(file))
        except Exception as e:
            logger.error(f"Error accessing directory {dir_path}: {e}")
            self._failed_files.add(str(dir_path))
    
    async def get_documents_by_keyword(self, keyword: str, index_type: str = "edarehoquqy") -> AsyncGenerator[Union[EdarehoquqyDocument, AraDocument], None]:
        """Get all documents for a specific keyword"""
        id_dirs = await self.get_id_dirs(keyword)
        for dir_path in id_dirs:
            if documents := await self.load_json_files(dir_path, index_type):
                for document in documents:
                    # Add to categorized IDs set
                    self._categorized_ids.add(document.id_edarehoquqy)
                    yield document
            else:
                self._failed_ids.add(dir_path.name)

    async def get_uncategorized_documents(self, index_type: str = "edarehoquqy") -> AsyncGenerator[Union[EdarehoquqyDocument, AraDocument], None]:
        """
        Get all documents that don't belong to any keyword category.
        
        Returns:
            AsyncGenerator of EdarehoquqyDocument objects
        """
        
        try:
            # Now get all existing IDs from the root directory
            existing_ids = await self.get_all_existing_ids(str(self.root_dir))
            logger.info(f"Existing IDs: {len(existing_ids)}")
            
            # Process each uncategorized document
            processed_count = 0
            async for doc in self.elasticsearch_service.get_missing_documents(existing_ids, index_type):
                try:
                    # Ensure we have a proper dictionary before converting to EdarehoquqyDocument
                    if not isinstance(doc, dict):
                        logger.warning(f"Skipping invalid document: {doc}")
                        continue
                        
                    # Make sure id_edarehoquqy exists in the document
                    if "id_edarehoquqy" not in doc:
                        logger.warning(f"Document missing ID field: {doc}")
                        continue
                    
                    # First convert to document object with whatever data we have
                    document = EdarehoquqyDocument(**doc)
                    
                    # Then try to get title and summary if needed - with retries
                    if not document.title or document.title == "":
                        # Try up to 3 times to get the title
                        for attempt in range(3):
                            try:
                                if title := await get_doc_title(self.openai_client, f"Question: {document.question}\nAnswer: {document.answer}"):
                                    document.title = title
                                    break
                            except Exception as e:
                                logger.warning(f"Error getting title for document {document.id_edarehoquqy} (attempt {attempt+1}/3): {e}")
                                if attempt < 2:  # Don't sleep after the last attempt
                                    await asyncio.sleep(2)  # Brief pause before retry
                    
                    if not document.summary or document.summary == "":
                        # Try up to 3 times to get the summary
                        for attempt in range(3):
                            try:
                                if summary := await get_each_doc_summary(self.openai_client, f"Question: {document.question}\nAnswer: {document.answer}"):
                                    document.summary = summary
                                    break
                            except Exception as e:
                                logger.warning(f"Error getting summary for document {document.id_edarehoquqy} (attempt {attempt+1}/3): {e}")
                                if attempt < 2:  # Don't sleep after the last attempt
                                    await asyncio.sleep(2)  # Brief pause before retry
                    
                    processed_count += 1
                    if processed_count % 100 == 0:
                        logger.info(f"Processed {processed_count} uncategorized documents so far")
                    
                    yield document
                except Exception as e:
                    logger.warning(f"Error processing uncategorized document: {e}")
                    continue

        except Exception as e:
            logger.error(f"Error loading uncategorized documents: {e}")
            import traceback
            logger.error(f"Traceback: {traceback.format_exc()}")
                
    @property
    def failed_ids(self) -> Set[str]:
        """Get the set of IDs that failed to load"""
        return self._failed_ids
        
    @property
    def failed_files(self) -> Set[str]:
        """Get the set of files that failed to load"""
        return self._failed_files
    



class WordExporter:
    def __init__(self, file_loader: Loader, output_dir: str = None, max_docs_per_file: int = 999):
        """
        Initialize a WordExporter to export documents to Word files.
        
        Args:
            file_loader: FileLoader instance to load documents from
            output_dir: Optional output directory to save Word documents to.
                       If not provided, documents will be saved in the keyword directories.
            max_docs_per_file: Maximum number of documents per Word file (default 999)
        """
        self.file_loader = file_loader
        self.output_dir = output_dir
        self.max_docs_per_file = max_docs_per_file
    
    def _get_new_composer(self, doc: Document) -> Composer:
        return Composer(doc)
    
    def _setup_styles(self, doc: Document) -> None:
        """Set up custom document styles for a professional appearance"""
        # Title style - large, bold, and prominent
        title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Arial'  # More impactful font
        title_font.size = Pt(18)   # Larger size for better visibility
        title_font.bold = True
        title_font.color.rgb = RGBColor(0, 0, 0)  # Black for maximum readability
        title_style.paragraph_format.space_after = Pt(16)  # More space after title
        title_style.paragraph_format.space_before = Pt(8)  # Add space before title
        title_style.paragraph_format.keep_with_next = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center alignment
        title_style.paragraph_format.border_bottom = True  # Add bottom border
        title_style.paragraph_format.border_bottom_color = RGBColor(0, 51, 102)  # Professional blue border
        title_style.paragraph_format.border_bottom_size = Pt(1)  # Border thickness
        
        # Metadata style - clean and distinct
        meta_style = doc.styles.add_style('MetadataLabel', WD_STYLE_TYPE.PARAGRAPH)
        meta_font = meta_style.font
        meta_font.name = 'Calibri'
        meta_font.size = Pt(12)
        meta_font.bold = True
        meta_font.color.rgb = RGBColor(89, 89, 89)  # Dark gray
        meta_style.paragraph_format.space_before = Pt(6)
        meta_style.paragraph_format.space_after = Pt(3)
        
        # Content style - readable and clean
        content_style = doc.styles.add_style('ContentText', WD_STYLE_TYPE.PARAGRAPH)
        content_font = content_style.font
        content_font.name = 'Calibri'
        content_font.size = Pt(15)
        content_style.paragraph_format.space_before = Pt(6)
        content_style.paragraph_format.space_after = Pt(6)
        content_style.paragraph_format.line_spacing = 1.15
        
        # Heading style for sections
        heading_style = doc.styles.add_style('SectionHeading', WD_STYLE_TYPE.PARAGRAPH)
        heading_font = heading_style.font
        heading_font.name = 'Calibri'
        heading_font.size = Pt(16)
        heading_font.bold = True
        heading_font.color.rgb = RGBColor(0, 51, 102)  # Match title color
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)
        heading_style.paragraph_format.keep_with_next = True

    def _create_document(self, data: EdarehoquqyDocument, idx: int) -> Document:
        doc = Document()
        self._setup_styles(doc)

        # Add document number with centered alignment
        doc_num_para = doc.add_paragraph(style='MetadataLabel')
        doc_num_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc_num_para.add_run(f"Document #{idx}")

        # Add title with proper styling
        title_text = data.title if data.title else "Untitled Document"
        title_para = doc.add_paragraph(style='CustomTitle')
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.add_run("Title: ").bold = True
        title_para.add_run(title_text)
        
        # Add metadata section
        metadata_section = doc.add_paragraph(style='MetadataLabel')
        metadata_section.add_run('Document Metadata\n').bold = True
        
        # Add ID info
        p = doc.add_paragraph(style='MetadataLabel')
        p.add_run(f"ID: ").bold = True
        p.add_run(data.id_edarehoquqy)
        
        # Format metadata nicely
        if hasattr(data, 'metadata') and data.metadata:
            for key, value in data.metadata.items():
                p = doc.add_paragraph(style='MetadataLabel')
                p.add_run(f"{key}: ").bold = True
                if isinstance(value, dict):
                    p.add_run(", ".join(f"{k}: {v}" for k, v in value.items()))
                else:
                    p.add_run(str(value))
        
        # Add a small separator
        separator = doc.add_paragraph(style='ContentText')
        separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
        separator.add_run('—' * 15)  # Smaller, more subtle separator
        
        # Add question
        if hasattr(data, 'question') and data.question:
            q_para = doc.add_paragraph(style='ContentText')
            q_para.add_run("Question: ").bold = True
            q_para.add_run(data.question)
            
        # Add answer
        if hasattr(data, 'answer') and data.answer:
            a_para = doc.add_paragraph(style='ContentText')
            a_para.add_run("Answer: ").bold = True
            a_para.add_run(data.answer)

        # If no question/answer, add content
        elif hasattr(data, 'content') and data.content:
            content_para = doc.add_paragraph(style='ContentText')
            content_para.add_run(data.content)
            
        # Add summary if available
        if hasattr(data, 'summary') and data.summary:
            # Small separator before summary
            sum_separator = doc.add_paragraph(style='ContentText')
            sum_separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sum_separator.add_run('—' * 15)
            
            doc.add_heading('Summary', level=2)
            summary_para = doc.add_paragraph(style='ContentText')
            summary_para.add_run(data.summary)
                
        # Add page break at the end
        doc.add_page_break()
        
        return doc
    
    def _create_ara_document(self, data: AraDocument, idx: int) -> Document:
        """
        Create a Word document specifically for ARA documents with professional formatting.
        
        Args:
            data: ARA document data to format
            idx: Index number for the document (for numbering)
            
        Returns:
            Document: Formatted Word document
        """
        try:
            doc = Document()
            
            # Set up styles
            self._setup_styles(doc)
            
            # Add document number with centered alignment
            doc_num_para = doc.add_paragraph(style='MetadataLabel')
            doc_num_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc_num_para.add_run(f"Document #{idx}")
            
            # Add title with proper styling
            title_text = data.title if hasattr(data, 'title') and data.title else "بدون عنوان"
            title_para = doc.add_paragraph(style='CustomTitle')
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.add_run("Title: ").bold = True
            title_para.add_run(title_text)
            
            # Add metadata section
            metadata_section = doc.add_paragraph(style='MetadataLabel')
            metadata_section.add_run('Document Metadata\n').bold = True
            
            # Add ID info
            p = doc.add_paragraph(style='MetadataLabel')
            p.add_run(f"ID: ").bold = True
            p.add_run(data.id_ara if hasattr(data, 'id_ara') else "Unknown")
            
            # Add metadata fields
            if hasattr(data, 'metadata_obj') and data.metadata_obj:
                metadata = data.metadata_obj
                
                metadata_fields = {
                    "authority_type": "نوع مرجع",
                    "judgment_type": "نوع حکم",
                    "judgment_group": "گروه حکم",
                    "branch": "شعبه",
                    "judge": "قاضی",
                    "session_number": "شماره جلسه",
                    "petition_date": "تاریخ دادخواست",
                    "final_judgment": "حکم نهایی",
                    "selected_document_judgment": "گزیده حکم"
                }
                
                for field_key, field_label in metadata_fields.items():
                    if field_key in metadata and metadata[field_key]:
                        p = doc.add_paragraph(style='MetadataLabel')
                        p.add_run(f"{field_label}: ").bold = True
                        value = metadata[field_key]
                        if isinstance(value, list):
                            p.add_run(", ".join(str(item) for item in value))
                        else:
                            p.add_run(str(value))
                
                # Add date information if available
                if hasattr(data, 'date') and data.date:
                    if hasattr(data.date, 'shamsi') and data.date.shamsi and data.date.shamsi != "0001/01/01":
                        p = doc.add_paragraph(style='MetadataLabel')
                        p.add_run("تاریخ شمسی: ").bold = True
                        p.add_run(data.date.shamsi)
                    
                    if hasattr(data.date, 'gregorian') and data.date.gregorian and data.date.gregorian != "0001/01/01":
                        p = doc.add_paragraph(style='MetadataLabel')
                        p.add_run("تاریخ میلادی: ").bold = True
                        p.add_run(data.date.gregorian)
            
            # Add a small separator
            separator = doc.add_paragraph(style='ContentText')
            separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
            separator.add_run('—' * 15)  # Smaller, more subtle separator
            
            # Add summary if available
            if hasattr(data, 'summary') and data.summary:
                doc.add_heading('Summary', level=2)
                summary_para = doc.add_paragraph(style='ContentText')
                summary_para.add_run(data.summary)
            
            # Add message section (specific to ARA documents)
            if hasattr(data, 'message') and data.message:
                doc.add_heading('Message', level=2)
                message_para = doc.add_paragraph(style='ContentText')
                message_para.add_run(data.message)
            
            # Add content section
            if hasattr(data, 'content') and data.content:
                doc.add_heading('Content', level=2)
                content_para = doc.add_paragraph(style='ContentText')
                content_para.add_run(data.content)
            
            # Add documents section if available
            if hasattr(data, 'documents') and data.documents:
                doc.add_heading('Documents', level=2)
                docs_para = doc.add_paragraph(style='ContentText')
                docs_para.add_run(data.documents)
            
            # Add page break at the end
            doc.add_page_break()
            
            return doc
        except Exception as e:
            logger.error(f"Error creating ARA document: {e}")
            import traceback
            logger.error(f"ARA document creation error details: {traceback.format_exc()}")
            
            # Create a minimal valid document to avoid breaking the process
            doc = Document()
            doc.add_paragraph(f"Error creating ARA document {idx}: {str(e)}")
            doc.add_page_break()
            return doc
    
    async def export_ara_keyword_docs(self, keyword: str) -> Optional[List[str]]:
        """Export all documents for a keyword to Word documents, splitting into multiple files if needed
        
        Args:
            keyword: The keyword to export documents for
            
        Returns:
            List of paths to the saved Word documents if successful, None otherwise
        """
        try:
            # Initialize variables
            output_paths = []
            doc_count = 0
            file_index = 1
            total_count = 0
            
            # Create first document
            master_doc = Document()
            self._setup_styles(master_doc)
            composer = Composer(master_doc)
            
            # Process each document for this keyword
            async for data in self.file_loader.get_documents_by_keyword(keyword, index_type="ara"):
                # Check if we need to start a new file
                if doc_count >= self.max_docs_per_file:
                    # Save the current file
                    if self.output_dir:
                        # Use specified output directory
                        output_path = os.path.join(self.output_dir, f"{keyword}_{file_index}.docx")
                    else:
                        # Save in the keyword directory
                        output_path = os.path.join(str(self.file_loader.root_dir), keyword, f"{keyword}_{file_index}.docx")  
                    
                    # Ensure directory exists
                    os.makedirs(os.path.dirname(output_path), exist_ok=True)
                    
                    # Save the document
                    composer.save(output_path)
                    output_paths.append(output_path)
                    logger.info(f"Saved document part {file_index} for '{keyword}' with {doc_count} documents to {output_path}")
                    
                    # Start a new document
                    master_doc = Document()
                    self._setup_styles(master_doc)
                    composer = Composer(master_doc)
                    doc_count = 0
                    file_index += 1
                
                # Add this document to the current file
                doc = self._create_ara_document(data, total_count + 1)
                composer.append(doc)
                doc_count += 1
                total_count += 1
            
            # Save the final document if it contains any documents
            if doc_count > 0:
                # Determine file name based on whether we already created other files
                file_suffix = f"_{file_index}" if file_index > 1 else ""
                
                if self.output_dir:
                    # Use specified output directory
                    output_path = os.path.join(self.output_dir, f"{keyword}{file_suffix}.docx")
                else:
                    # Save in the keyword directory
                    output_path = os.path.join(str(self.file_loader.root_dir), keyword, f"{keyword}{file_suffix}.docx")  
                    
                # Ensure directory exists
                os.makedirs(os.path.dirname(output_path), exist_ok=True)
                
                # Save the document
                composer.save(output_path)
                output_paths.append(output_path)
                logger.info(f"Saved final document part for '{keyword}' with {doc_count} documents to {output_path}")
            
            if not output_paths:
                logger.warning(f"No documents found for keyword: {keyword}")
                return None
                
            logger.info(f"Successfully exported {total_count} documents for '{keyword}' across {len(output_paths)} files")
            return output_paths
            
        except Exception as e:
            logger.error(f"Error creating Word document for keyword '{keyword}': {e}")
            import traceback
            logger.error(f"Error details: {traceback.format_exc()}")
            return None
        
    async def export_keyword_to_word(self, keyword: str) -> Optional[List[str]]:
        """Export all documents for a keyword to Word documents, splitting into multiple files if needed
        
        Args:
            keyword: The keyword to export documents for
            
        Returns:
            List of paths to the saved Word documents if successful, None otherwise
        """
        try:
            # Initialize variables
            output_paths = []
            doc_count = 0
            file_index = 1
            total_count = 0
            
            # Create first document
            master_doc = Document()
            self._setup_styles(master_doc)
            composer = Composer(master_doc)
            
            # Process each document for this keyword
            async for data in self.file_loader.get_documents_by_keyword(keyword):
                # Check if we need to start a new file
                if doc_count >= self.max_docs_per_file:
                    # Save the current file
                    if self.output_dir:
                        # Use specified output directory
                        output_path = os.path.join(self.output_dir, f"{keyword}_{file_index}.docx")
                    else:
                        # Save in the keyword directory
                        output_path = os.path.join(str(self.file_loader.root_dir), keyword, f"{keyword}_{file_index}.docx")
                    
                    # Ensure directory exists
                    os.makedirs(os.path.dirname(output_path), exist_ok=True)
                    
                    # Save the document
                    composer.save(output_path)
                    output_paths.append(output_path)
                    logger.info(f"Saved document part {file_index} for '{keyword}' with {doc_count} documents to {output_path}")
                    
                    # Start a new document
                    master_doc = Document()
                    self._setup_styles(master_doc)
                    composer = Composer(master_doc)
                    doc_count = 0
                    file_index += 1
                
                # Add this document to the current file
                doc = self._create_document(data, total_count + 1)
                composer.append(doc)
                doc_count += 1
                total_count += 1
            
            # Skip if no documents were found
            if total_count == 0:
                logger.warning(f"No documents found for keyword: {keyword}")
                return None
                
            # Save the final document if it contains any documents
            if doc_count > 0:
                # Determine file name based on whether we already created other files
                file_suffix = f"_{file_index}" if file_index > 1 else ""
                
                if self.output_dir:
                    # Use specified output directory
                    output_path = os.path.join(self.output_dir, f"{keyword}{file_suffix}.docx")
                else:
                    # Save in the keyword directory
                    output_path = os.path.join(str(self.file_loader.root_dir), keyword, f"{keyword}{file_suffix}.docx")
                    
                # Ensure directory exists
                os.makedirs(os.path.dirname(output_path), exist_ok=True)
                
                # Save the document
                composer.save(output_path)
                output_paths.append(output_path)
                logger.info(f"Saved final document part for '{keyword}' with {doc_count} documents to {output_path}")
            
            logger.info(f"Successfully exported {total_count} documents for '{keyword}' across {len(output_paths)} files")
            return output_paths
            
        except Exception as e:
            logger.error(f"Error creating Word document for keyword '{keyword}': {e}")
            import traceback
            logger.error(f"Error details: {traceback.format_exc()}")
            return None

    async def export_uncategorized_documents(self, index_type: str = "edarehoquqy") -> Optional[List[str]]:
        """
        Export all uncategorized documents to Word documents, splitting into multiple files if needed,
        and also create HTML files
        
        Args:
            index_type: Type of index to process ('edarehoquqy' or 'ara')
            
        Returns:
            List of paths to the saved Word documents if successful, None otherwise
        """
        try:
            logger.info(f"Exporting uncategorized documents for index type: {index_type}")
            
            # Initialize variables
            output_paths = []
            doc_count = 0
            file_index = 1
            total_count = 0
            
            # Create first document
            master_doc = Document()
            self._setup_styles(master_doc)
            composer = Composer(master_doc)
            
            # The uncategorized keyword
            uncategorized_keyword = "بدون-دسته-بندی"
            
            # Make the directory for JSON/HTML files
            uncategorized_dir = os.path.join(str(self.file_loader.root_dir), uncategorized_keyword)
            os.makedirs(uncategorized_dir, exist_ok=True)
            
            # Process each uncategorized document
            async for data in self.file_loader.get_uncategorized_documents(index_type):
                try:
                    # Check if we need to start a new file
                    if doc_count >= self.max_docs_per_file:
                        # Save the current file
                        if self.output_dir:
                            # Use specified output directory
                            output_path = os.path.join(self.output_dir, f"{uncategorized_keyword}_{file_index}.docx")
                        else:
                            # Save in the uncategorized directory
                            output_path = os.path.join(uncategorized_dir, f"{uncategorized_keyword}_{file_index}.docx")
                        
                        # Ensure directory exists
                        os.makedirs(os.path.dirname(output_path), exist_ok=True)
                        
                        # Save the document
                        composer.save(output_path)
                        output_paths.append(output_path)
                        logger.info(f"Saved document part {file_index} for uncategorized with {doc_count} documents to {output_path}")
                        
                        # Start a new document
                        master_doc = Document()
                        self._setup_styles(master_doc)
                        composer = Composer(master_doc)
                        doc_count = 0
                        file_index += 1
                    
                    # Determine document type based on attributes
                    is_ara_document = hasattr(data, 'id_ara') and not hasattr(data, 'id_edarehoquqy')
                    
                    # Add to Word document using appropriate method
                    if index_type == "ara" or is_ara_document:
                        doc = self._create_ara_document(data, total_count + 1)
                        doc_id = data.id_ara
                    else:
                        doc = self._create_document(data, total_count + 1)
                        doc_id = data.id_edarehoquqy
                    
                    composer.append(doc)
                    doc_count += 1
                    total_count += 1
                    
                    # Create directory for this document's files
                    doc_dir = AsyncPath(uncategorized_dir) / doc_id
                    await doc_dir.mkdir(exist_ok=True)
                    
                    # Save JSON file
                    json_path = doc_dir / f"{doc_id}.json"
                    async with aiofiles.open(json_path, mode="w", encoding="utf-8") as file:
                        await file.write(json.dumps(data.model_dump(), ensure_ascii=False, indent=4))
                        
                    # Save HTML file
                    try:
                        await save_html_file(data, doc_dir)
                        logger.info(f"Saved HTML file for document {doc_id}")
                    except Exception as e:
                        logger.error(f"Error saving HTML file for {doc_id}: {e}")
                except Exception as e:
                    logger.error(f"Error processing uncategorized document: {e}")
                    import traceback
                    logger.error(f"Error details: {traceback.format_exc()}")
                    continue
            
            # Skip if no documents were found
            if total_count == 0:
                logger.info("No uncategorized documents found")
                return None
            
            # Save the final document if it contains any documents
            if doc_count > 0:
                # Determine file name based on whether we already created other files
                file_suffix = f"_{file_index}" if file_index > 1 else ""
                
                if self.output_dir:
                    # Use specified output directory
                    output_path = os.path.join(self.output_dir, f"{uncategorized_keyword}{file_suffix}.docx")
                else:
                    # Save in the uncategorized directory
                    output_path = os.path.join(uncategorized_dir, f"{uncategorized_keyword}{file_suffix}.docx")
                    
                # Ensure directory exists
                os.makedirs(os.path.dirname(output_path), exist_ok=True)
                
                # Save the document
                composer.save(output_path)
                output_paths.append(output_path)
                logger.info(f"Saved final document part for uncategorized with {doc_count} documents to {output_path}")
            
            logger.info(f"Successfully exported {total_count} uncategorized documents across {len(output_paths)} files")
            return output_paths
                
        except Exception as e:
            logger.error(f"Error creating Word document for uncategorized documents: {e}")
            import traceback
            logger.error(f"Error details: {traceback.format_exc()}")
            return None
    
    async def export_all_keywords(self, limit: int = None, index_type: str = "edarehoquqy") -> Dict[str, List[str]]:
        """Export all keywords to separate Word documents, splitting into multiple files if needed
        
        Args:
            limit: Optional limit on the number of keywords to process
            index_type: Type of index to process ('edarehoquqy' or 'ara')
            
        Returns:
            Dictionary mapping keywords to lists of saved Word document paths
        """
        results = {}
        logger.info(f"Starting export of keywords with index type: {index_type}")
        
        try:
            # Get all keywords
            keywords = await self.file_loader.get_keywords()
            logger.info(f"Found {len(keywords)} keywords to process")
            
            # Apply limit if provided
            if limit:
                logger.info(f"Limiting to first {limit} keywords")
                keywords = keywords[:limit]
            
            # Process each keyword and save its document
            for i, keyword in enumerate(keywords):
                logger.info(f"Processing keyword {i+1}/{len(keywords)}: '{keyword}'")
                if index_type == "ara":
                    output_paths = await self.export_ara_keyword_docs(keyword)
                else:
                    output_paths = await self.export_keyword_to_word(keyword)
                
                if output_paths:
                    logger.info(f"Successfully exported keyword '{keyword}' to {len(output_paths)} files")
                    results[keyword] = output_paths
                else:
                    logger.warning(f"Failed to export keyword '{keyword}'")
            
            logger.info(f"Successfully exported {len(results)} of {len(keywords)} keywords to Word documents")
            return results
            
        except Exception as e:
            logger.error(f"Error in export_all_keywords: {e}")
            import traceback
            logger.error(f"Error details: {traceback.format_exc()}")
            return results  # Return whatever results we have so far

    async def export_documents_by_keyword(self, keyword: str, index_type: str = "edarehoquqy") -> Optional[List[str]]:
        """Export all documents for a specific keyword to Word documents
        
        Args:
            keyword: The keyword to export documents for
            index_type: Type of index to process ('edarehoquqy' or 'ara')
            
        Returns:
            List of paths to the saved Word documents if successful, None otherwise
        """
        logger.info(f"Exporting documents for keyword '{keyword}' with index type '{index_type}'")
        
        try:
            if index_type == "ara":
                return await self.export_ara_keyword_docs(keyword)
            else:
                return await self.export_keyword_to_word(keyword)
        except Exception as e:
            logger.error(f"Error exporting documents for keyword '{keyword}': {e}")
            import traceback
            logger.error(f"Error details: {traceback.format_exc()}")
            return None


async def main():
    """Main function to export documents to Word files"""
    # Initialize the loader with the search directory
    hosts = [os.getenv("ELASTICSEARCH_HOST")]
    username = os.getenv("ELASTICSEARCH_USERNAME")
    password = os.getenv("ELASTICSEARCH_PASSWORD")
    index_name = os.getenv("ELASTICSEARCH_INDEX_NAME")
    openai_client = AsyncOpenAI(api_key=os.getenv("OPENAI_API_KEY"), base_url=os.getenv("OPENAI_BASE_URL"))
    logger.info(f"Hosts: {hosts}, Username: {username}, Password: {password}, Index name: {index_name}")
    
    # Define search directory in current working directory
    # search_dir = os.path.join(os.getcwd(), "search")
    # Use absolute path instead
    search_dir = "/home/sadeghian/edarehoquqy/html-outputls/search"
    
    # Create the directory if it doesn't exist
    if not os.path.exists(search_dir):
        logger.info(f"Creating search directory at {search_dir}")
        os.makedirs(search_dir, exist_ok=True)
    else:
        logger.info(f"Using existing search directory at {search_dir}")
    
    # Use ElasticsearchService with async context manager
    async with ElasticsearchService(hosts=hosts, username=username, password=password, index_name=index_name) as es:
        await es.ensure_index_exists()
        loader = Loader(search_dir, es, openai_client)
        exporter = WordExporter(loader)
            
        # Export all keywords to Word documents
        # results = await exporter.export_all_keywords()
        
        # Log results
        # if results:
        #     logger.info(f"Successfully exported {len(results)} keywords to Word documents:")
        #     for keyword, paths in results.items():
        #         logger.info(f"  - {keyword}: {paths}")
        # else:
        #     logger.warning("No documents were exported")
        
        # Now handle uncategorized documents
        logger.info("Now processing uncategorized documents...")
        uncategorized_paths = await exporter.export_uncategorized_documents()
        if uncategorized_paths:
            logger.info(f"Successfully exported uncategorized documents to: {uncategorized_paths}")
        else:
            logger.info("No uncategorized documents found")
            
        async with aiofiles.open("failed_status.txt", mode="w", encoding="utf-8") as file:
            # Log failures
            failed_ids = loader.failed_ids
            if failed_ids:
                logger.warning(f"Failed to load {len(failed_ids)} document IDs")
                await file.write(f"Failed to load {len(failed_ids)} document IDs: {failed_ids}\n")
            
            failed_files = loader.failed_files
            if failed_files:
                logger.warning(f"Failed to load {len(failed_files)} files")
                await file.write(f"Failed to load {len(failed_files)} files: {failed_files}\n")



if __name__ == "__main__":
    asyncio.run(main())


    
    

